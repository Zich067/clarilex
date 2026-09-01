"""組裝論文 DOCX,符合國立臺中科技大學 114 年學位論文格式規範。

對應規範要點:
  - A4,邊界全部 3 cm,頁尾頁碼置中、距底 2 cm、10 pt
  - 中文新細明體、英文 Times New Roman、內文 12 pt、單行間距
  - 第一層標題(章) 新細明體粗體 16 pt
  - 第二層標題(節) 新細明體粗體 14 pt
  - 第三層標題       新細明體粗體 12 pt
  - 摘要至圖目次:小寫羅馬數字 i, ii, iii
  - 論文本文至附錄:阿拉伯數字 1, 2, 3
  - 編印順序: 封面 → 書名頁 → 中摘 → 英摘 → 誌謝(可略) → 目次 → 表目次 → 圖目次
              → 本文 → 參考文獻 → 附錄(可略)

執行:
    python scripts/build_thesis_docx.py
輸出:
    thesis/論文_許紫晴.docx
"""

from __future__ import annotations

import hashlib
import re
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = ROOT.parent              # /Users/ching/Desktop/thesis-system
THESIS_DIR = PROJECT_ROOT / "thesis"
FIG_DIR = THESIS_DIR / "figures"
CODE_FIG_DIR = FIG_DIR / "code"         # Carbon 程式碼圖快取
OUT_PATH = THESIS_DIR / "論文_許紫晴.docx"
WATERMARK_PATH = ROOT / "data" / "watermark_nutc.jpeg"
APPROVAL_PATH = ROOT / "data" / "approval_hsu.jpg"   # 學位考試委員會審定書掃描影像

MATH_FIG_DIR = FIG_DIR / "math"          # 數學式圖快取

sys.path.insert(0, str(Path(__file__).resolve().parent))
from make_code_carbon import render_carbon  # noqa: E402
from make_math_png import render_math  # noqa: E402


# ─── meta ───
META = {
    "dept_zh": "國立臺中科技大學資訊工程系碩士班",
    "dept_en": "Department of Computer Science and Information Engineering",
    "university_en": "National Taichung University of Science and Technology",
    "degree_en": "Master of Engineering",
    "title_zh": "基於檢索增強生成技術之租賃買賣法律文件智慧分析與風險評估系統",
    "title_en": (
        "An Intelligent Lease and Sale Legal Document Analysis and Risk "
        "Assessment System Based on Retrieval-Augmented Generation"
    ),
    "advisors": [
        {"zh": "張家瑋", "en": "Jia-Wei Chang"},
        {"zh": "馬豪尚", "en": "Hao-Shang Ma"},
    ],
    "student_zh": "許紫晴",
    "student_en": "Tzu-Ching Hsu",
    "year_roc": "115",
    "month_zh": "7",
    "month_en": "July 2026",
}


# ────────────────────────── 字型 helper ──────────────────────────


def _set_run_fonts(
    run,
    cjk: str = "新細明體",
    ascii_: str = "Times New Roman",
    size_pt: float | None = None,
    bold: bool = False,
) -> None:
    """同時設定 CJK 與 ASCII 字型,顯式寫入 rPr/rFonts/eastAsia 標籤。"""
    run.font.name = ascii_
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_)
    rfonts.set(qn("w:hAnsi"), ascii_)
    rfonts.set(qn("w:eastAsia"), cjk)
    rfonts.set(qn("w:cs"), ascii_)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold:
        run.font.bold = True


def add_centered_paragraph(
    doc: DocxDocument,
    text: str,
    *,
    size_pt: float = 12,
    bold: bool = False,
    cjk: str = "新細明體",
    space_before: float = 0,
    space_after: float = 0,
) -> Paragraph:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    _set_run_fonts(r, cjk=cjk, size_pt=size_pt, bold=bold)
    return p


def add_left_paragraph(
    doc: DocxDocument,
    text: str = "",
    *,
    size_pt: float = 12,
    bold: bool = False,
    cjk: str = "新細明體",
    indent: float = 0,
    space_before: float = 0,
    space_after: float = 0,
) -> Paragraph:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    if indent:
        pf.left_indent = Cm(indent)
    if text:
        r = p.add_run(text)
        _set_run_fonts(r, cjk=cjk, size_pt=size_pt, bold=bold)
    return p


def add_stacked_left_block(
    doc: DocxDocument,
    lines: "list[str]",
    *,
    size_pt: float = 14,
    bold: bool = False,
    cjk: str = "新細明體",
    indent: float = 0,
) -> Paragraph:
    """左對齊多行掛名區塊(同段內以換行串接)。

    掛名採「左對齊 + 整塊縮排視覺置中」而非逐行置中:逐行置中會把續行前導之全形
    空格連同名字一起重新置中、使空格失效(第二位指導教授飄移);左對齊時前導空格
    方能生效,使「馬豪尚」精準對齊至「張家瑋」正下方。indent 控整塊水平位置。
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    if indent:
        pf.left_indent = Cm(indent)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()  # 段內軟換行(非新段落),使各行共用同一 left_indent
        r = p.add_run(line)
        _set_run_fonts(r, cjk=cjk, size_pt=size_pt, bold=bold)
    return p


def tighten_toc_styles(doc: DocxDocument) -> None:
    """壓緊自動目次(目次／表目次／圖目次)之樣式 TOC 1–4:單行間距、無段前後距。
    Word 預設 TOC 樣式帶較大段後距,使目次過長、末項易孤立溢頁;此處壓緊即可。"""
    _sizes = {"TOC 1": 13, "TOC 2": 11, "TOC 3": 11, "TOC 4": 11}  # 規範:章 13 級、節/子節 11 級
    for name, _sz in _sizes.items():
        try:
            st = doc.styles[name]
        except KeyError:
            try:
                st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            except Exception:  # noqa: BLE001
                continue
        st.font.size = Pt(_sz)
        _rpr = st.element.get_or_add_rPr()
        _rf = _rpr.find(qn("w:rFonts"))
        if _rf is None:
            _rf = OxmlElement("w:rFonts")
            _rpr.append(_rf)
        _rf.set(qn("w:ascii"), "Times New Roman")
        _rf.set(qn("w:hAnsi"), "Times New Roman")
        _rf.set(qn("w:eastAsia"), "新細明體")
        pf = st.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0  # 規範:單行間距


def add_page_break(doc: DocxDocument) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_blank_line(doc: DocxDocument, count: int = 1, size_pt: float = 12) -> None:
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run("")
        _set_run_fonts(r, size_pt=size_pt)


# ────────────────────────── 頁碼 helper ──────────────────────────


def _make_page_number_field(run) -> None:
    """在 run 中插入 PAGE 欄位 (Word 開檔時會自動更新)。"""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE   \\* MERGEFORMAT"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)
    run._element.append(t)
    run._element.append(fld_end)


def _setup_section_page_format(
    section,
    *,
    margin_cm: float = 3,
    page_num_fmt: str = "decimal",
    start: int | None = 1,
    restart: bool = False,
) -> None:
    """設定 section 頁面大小、邊界、頁碼格式。

    page_num_fmt: "decimal" (1,2,3) 或 "lowerRoman" (i,ii,iii)
    restart: True 時插入 type="restart" 重置頁碼
    """
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(2)

    # 設定 pgNumType
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), page_num_fmt)
    if restart and start is not None:
        pgNumType.set(qn("w:start"), str(start))
    elif start is not None and start != 1:
        pgNumType.set(qn("w:start"), str(start))


def _add_footer_page_number(section) -> None:
    """頁尾置中放頁碼,10 pt 新細明體。"""
    footer = section.footer
    footer.is_linked_to_previous = False
    # 清掉預設段落並重建
    for p in list(footer.paragraphs):
        p_elem = p._element
        p_elem.getparent().remove(p_elem)
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    r = p.add_run()
    _set_run_fonts(r, size_pt=10)
    _make_page_number_field(r)


def add_section_break(doc: DocxDocument) -> None:
    """加一個 NEW_PAGE section break。"""
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    return new_section


# ────────────────────────── Markdown 解析 ──────────────────────────


_CODE_FENCE = re.compile(r"^```")
_DISPLAY_MATH = re.compile(r"^\s*\$\$(.+?)\$\$\s*$")
_HEADING = re.compile(r"^(#{1,6})\s+(.*)")
_LIST_BULLET = re.compile(r"^\s*[-*]\s+(.*)")
_LIST_NUMBERED = re.compile(r"^\s*(\d+)\.\s+(.*)")
# 縮排字母子項 (a. b. c. …)：須有前導空白,避免誤抓正文;單一小寫字母
_LIST_LETTER = re.compile(r"^\s+([a-z])\.\s+(.*)")
_BLOCKQUOTE = re.compile(r"^>\s?(.*)")
_TABLE_SEP = re.compile(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")
_FRONTMATTER = re.compile(r"^---\s*$")
_IMG = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_INLINE_CODE = re.compile(r"`([^`]+)`")
_BOLD = re.compile(r"\*\*([^*]+)\*\*")
_ITALIC = re.compile(r"\*([^*]+)\*")
_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_INLINE_MATH = re.compile(r"\$([^$\n]+?)\$")

# 行內數學降級對照（pandoc OMML 上線前之過渡；display $$ 仍走 PNG 不變）
_LATEX_SYMBOLS = [
    ("\\varnothing", "∅"), ("\\emptyset", "∅"),
    ("\\cup", "∪"), ("\\cap", "∩"), ("\\setminus", "∖"),
    ("\\neq", "≠"), ("\\leq", "≤"), ("\\geq", "≥"),
    ("\\cdot", "·"), ("\\times", "×"), ("\\in", "∈"),
    ("\\phi", "φ"), ("\\lambda", "λ"),
    ("\\log", "log"), ("\\min", "min"), ("\\max", "max"),
    ("\\cos", "cos"), ("\\sum", "Σ"),
    ("\\,", " "), ("\\;", " "), ("\\quad", "  "),
]


def _latex_to_text(s: str) -> str:
    """把行內 $...$ 之 LaTeX 降級為可讀 Unicode 純文字。

    過渡用:現行 matplotlib mathtext 無法處理 \\mathcal / \\mathbb 等符號、
    且不渲染行內式,故先轉可讀文字以止住裸 $ 漏字;待 pandoc OMML 再改原生方程式。
    """
    s = re.sub(r"\\(?:mathrm|mathcal|mathbb|operatorname|text|mathbf)\{([^{}]*)\}", r"\1", s)
    for k, v in _LATEX_SYMBOLS:
        s = s.replace(k, v)
    s = re.sub(r"\\[a-zA-Z]+", "", s)
    s = s.replace("{", "").replace("}", "").replace("\\", "")
    return re.sub(r"[ \t]+", " ", s).strip()


def render_inline(text: str) -> list[tuple[str, dict]]:
    """把一行文字拆成 [(plain, attrs), ...] 的 segments。
    attrs: {bold, code, italic}
    這版本實作簡化:僅處理粗體與行內 code,其他原樣輸出。
    """
    segs: list[tuple[str, dict]] = []
    i = 0
    while i < len(text):
        # bold
        m = _BOLD.match(text, i)
        if m:
            segs.append((m.group(1), {"bold": True}))
            i = m.end()
            continue
        # italic(單星號 *...*;放在粗體之後,確保 ** 先被吃掉)
        m = _ITALIC.match(text, i)
        if m:
            segs.append((m.group(1), {"italic": True}))
            i = m.end()
            continue
        # inline code
        m = _INLINE_CODE.match(text, i)
        if m:
            segs.append((m.group(1), {"code": True}))
            i = m.end()
            continue
        # link → 顯示文字部分
        m = _LINK.match(text, i)
        if m:
            segs.append((m.group(1), {}))
            i = m.end()
            continue
        # 行內數學 $...$ → 標記為 math segment,由 emit 端轉成 Word 原生方程式(OMML)
        m = _INLINE_MATH.match(text, i)
        if m:
            segs.append((m.group(1), {"math": True}))
            i = m.end()
            continue
        # plain: walk to next special
        next_idx = len(text)
        for pat in (_BOLD, _ITALIC, _INLINE_CODE, _LINK, _INLINE_MATH):
            m = pat.search(text, i)
            if m and m.start() < next_idx:
                next_idx = m.start()
        segs.append((text[i:next_idx], {}))
        i = next_idx
    return segs


def _emit_inline(
    p,
    text: str,
    *,
    size_pt: float = 12,
    code_size: float | None = None,
    code_ascii: str = "Consolas",
) -> None:
    """把一行文字(粗體 / 行內 code / 行內公式)emit 到段落 p。

    行內 $...$ 轉成 Word 原生方程式(OMML <m:oMath>);pandoc 不可用時
    才退回可讀純文字(絕不留下裸 $)。
    """
    for seg, attrs in render_inline(text):
        if not seg:
            continue
        if attrs.get("math"):
            oms = _latex_to_omath(seg, display=False)
            if oms:
                for om in oms:
                    p._p.append(om)
                continue
            r = p.add_run(_latex_to_text(seg))
            _set_run_fonts(r, size_pt=size_pt)
            continue
        r = p.add_run(seg)
        if attrs.get("code"):
            _set_run_fonts(
                r,
                cjk="新細明體",
                ascii_=code_ascii,
                size_pt=(code_size if code_size is not None else size_pt - 1),
            )
        else:
            _set_run_fonts(r, size_pt=size_pt, bold=attrs.get("bold", False))
            if attrs.get("italic"):
                r.italic = True


def emit_paragraph(doc: DocxDocument, text: str, *, size_pt: float = 12) -> None:
    """內文段落:新細明體 12 級、單行間距、首行縮排 2 字元、段距 0。

    依 NUTC 附件十/十一:對齊靠左、行距單行、首行縮排 2 字元。
    """
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(size_pt * 2)  # 首行縮排 2 字元
    _emit_inline(p, text, size_pt=size_pt)
    return p


def emit_heading(doc: DocxDocument, level: int, text: str,
                 page_break_before: bool = False) -> None:
    """level 1 = 章 16pt; 2 = 節 14pt; 3 = 12pt 粗;
    page_break_before=True 時,本標題前換頁(由段落屬性達成,不留懸空空段落,
    避免前一章寫到頁尾時 add_page_break 溢出成整頁空白)。"""
    text = text.strip()
    if level == 1:
        size = 16
        space_before = 18
        space_after = 12
    elif level == 2:
        size = 14
        space_before = 12
        space_after = 8
    else:
        size = 12
        space_before = 8
        space_after = 4
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if page_break_before:
        pf.page_break_before = True
    # 規範(附件十一)僅定義三層「編號」標題。
    # 三層以內且為編號標題 → 設 Heading style 進目次;
    # 無編號之三層標題(Phase/Case 等)與第四層以上 → 呈現為粗體導引句,不進目次。
    numbered = bool(re.match(r"\d", text))
    if level <= 2 or (level == 3 and numbered):
        p.style = f"Heading {level}"
    r = p.add_run(text)
    _set_run_fonts(r, size_pt=size, bold=True)


def emit_list_item(
    doc: DocxDocument,
    text: str,
    *,
    numbered: bool = False,
    number: str | None = None,
    indent_base: float = 0.74,
) -> None:
    if numbered and number is not None:
        # 以 markdown 原始編號當文字輸出,避免 Word「List Number」全文件連號,
        # 並保住「第 N 點」之交叉引用正確。懸掛縮排使換行對齊。
        # indent_base 控整塊左縮排(字母子項傳更大值,視覺上巢狀於數字項之下)。
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_after = Pt(0)
        pf.left_indent = Cm(indent_base)
        pf.first_line_indent = Cm(-0.74)
        r = p.add_run(f"{number}. ")
        _set_run_fonts(r, size_pt=12)
        _emit_inline(p, text, size_pt=12, code_size=11)
        return
    style = "List Number" if numbered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(0)
    _emit_inline(p, text, size_pt=12, code_size=11)


def emit_code_block(doc: DocxDocument, lines: list[str], lang: str = "") -> None:
    """以等寬字型把程式碼當「可選取純文字」輸出於淺灰底單格表格內。

    碩論之程式碼應為可選取、可編輯之文字,而非深色截圖。故不再貼 Carbon PNG,
    改以 Consolas(英數)+ 新細明體(中文註解)10 級單行間距呈現,並保留縮排。
    """
    # 去除前後空行
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    if not lines:
        return

    _XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    # 淺灰底
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    tcPr.append(shd)

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    for idx, line in enumerate(lines):
        if idx > 0:
            para.add_run().add_break()
        run = para.add_run(line if line.strip() else " ")
        _set_run_fonts(run, cjk="新細明體", ascii_="Consolas", size_pt=10)
        # 保留行首縮排空白
        for t in run._element.findall(qn("w:t")):
            t.set(_XML_SPACE, "preserve")


_MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _latex_to_omath(latex: str, display: bool = True):
    """用 pandoc 把 LaTeX 轉為 Word 原生方程式(OMML <m:oMath>)元素清單;失敗回 None。"""
    import io
    import subprocess
    import zipfile
    from copy import deepcopy

    from lxml import etree

    md = f"$$\n{latex}\n$$\n" if display else f"${latex}$"
    try:
        out = subprocess.run(
            ["pandoc", "-f", "markdown+tex_math_dollars", "-t", "docx", "-o", "-"],
            input=md.encode("utf-8"), capture_output=True, check=True,
        ).stdout
        with zipfile.ZipFile(io.BytesIO(out)) as z:
            xml = z.read("word/document.xml")
        root = etree.fromstring(xml)
        oms = root.findall(f".//{{{_MATH_NS}}}oMath")
        return [deepcopy(o) for o in oms] or None
    except Exception as e:  # noqa: BLE001
        import sys
        print(
            f"[!] OMML 公式轉換失敗,該式將退回備援(latex={latex[:50]!r}): {e}",
            file=sys.stderr,
        )
        return None


_EQ_NO = [0]  # display 公式流水號(全文連續編號)


def emit_math_block(doc: DocxDocument, latex: str) -> None:
    """把 $$...$$ 顯示公式以 Word 原生方程式(OMML)置中插入,右側加全文連續編號 (n);
    pandoc 不可用時退回 PNG。版面:左tab→置中tab(公式)→右tab→(n)靠右對齊。"""
    _EQ_NO[0] += 1
    n = _EQ_NO[0]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    # 置中 tab(版面中央 7.5cm)+ 靠右 tab(右邊界 15cm = A4 21cm 減左右各 3cm)
    pf.tab_stops.add_tab_stop(Cm(7.5), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT)
    lead = p.add_run("\t")
    _set_run_fonts(lead, size_pt=12)

    oms = _latex_to_omath(latex, display=True)
    if oms:
        for om in oms:
            p._p.append(om)
    else:
        # ── fallback:pandoc 不可用時退回原 Computer Modern PNG 行為 ──
        digest = hashlib.md5(latex.encode("utf-8")).hexdigest()[:12]
        png = MATH_FIG_DIR / f"eq_{digest}.png"
        if not png.exists():
            render_math(latex, png)
        from PIL import Image

        with Image.open(png) as im:
            w_px, _ = im.size
        width_cm = min(11.0, max(3.0, w_px / 150.0))
        p.add_run().add_picture(str(png), width=Cm(width_cm))

    num = p.add_run(f"\t({n})")
    _set_run_fonts(num, size_pt=12)


def _border_el(tag: str, val: str, sz: int) -> "OxmlElement":
    e = OxmlElement(f"w:{tag}")
    e.set(qn("w:val"), val)
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), "000000")
    return e


def _apply_three_line_table(table) -> None:
    """套用台灣碩博論文標準三線表:僅頂線、表頭下分隔線、底線;

    無垂直框線、無內部橫線、無底色。線粗:頂/底 1.5 pt(sz=12)、
    表頭下分隔線 0.75 pt(sz=6)。對應規範:表格不上色、不加格線。
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    # 移除既有表格樣式參照(避免 Light Grid 等帶色帶框樣式)
    for ts in tblPr.findall(qn("w:tblStyle")):
        tblPr.remove(ts)
    # 重設表格層級邊框:只留頂、底
    for b in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(b)
    tbl_borders = OxmlElement("w:tblBorders")
    tbl_borders.append(_border_el("top", "single", 12))
    tbl_borders.append(_border_el("bottom", "single", 12))
    tbl_borders.append(_border_el("left", "none", 0))
    tbl_borders.append(_border_el("right", "none", 0))
    tbl_borders.append(_border_el("insideH", "none", 0))
    tbl_borders.append(_border_el("insideV", "none", 0))
    tblPr.append(tbl_borders)
    # 表頭列加下分隔線(三線表的中線),並清除所有儲存格底色
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for shd in tcPr.findall(qn("w:shd")):
                tcPr.remove(shd)
            if ri == 0:
                for tb in tcPr.findall(qn("w:tcBorders")):
                    tcPr.remove(tb)
                tc_borders = OxmlElement("w:tcBorders")
                tc_borders.append(_border_el("bottom", "single", 6))
                tcPr.append(tc_borders)


def _cell_wlen(text: str) -> int:
    """估算顯示寬度:中文/全形算 2、其餘算 1。數學式/行內碼以渲染後寬度估(視為短),
    避免 LaTeX 原始碼字數過長,使純符號欄被誤判為長文字欄而靠左。"""
    t = text or ""
    t = re.sub(r"\$[^$]*\$", "MM", t)   # $math$ 視為約 2 字寬
    t = re.sub(r"`[^`]*`", "MM", t)        # `code` 視為約 2 字寬
    t = re.sub(r"\*\*", "", t)
    return sum(2 if ord(c) > 0x2E80 else 1 for c in t)


def emit_table(doc: DocxDocument, headers: list[str], rows: list[list[str]]) -> None:
    if not headers:
        return
    ncol = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=ncol)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.allow_autofit = False
    _tblPr = t._tbl.tblPr
    _mar = OxmlElement("w:tblCellMar")
    for _side in ("left", "right"):
        _e = OxmlElement(f"w:{_side}")
        _e.set(qn("w:w"), "40")
        _e.set(qn("w:type"), "dxa")
        _mar.append(_e)
    _tblPr.append(_mar)
    # 依各欄最長內容估相對寬度,分配於 15.6cm(直向可用寬),使長文字換行、短欄不過寬。
    # 表頭以「半寬」計入(表頭可折兩行,不讓長表頭如「對應困難點（§1.2）」獨佔寬度),
    # 與內容最長者取大,作為相對寬度量。
    import math as _math
    content_max = []
    head_w = []
    for ci in range(ncol):
        cm = 1
        for row in rows:
            if ci < len(row):
                cm = max(cm, _cell_wlen(row[ci]))
        content_max.append(cm)
        head_w.append(_cell_wlen(headers[ci]))
    measure = [max(content_max[ci], _math.ceil(head_w[ci] / 2), 4) for ci in range(ncol)]
    # 抑制超長欄獨佔寬度(以欄平均之 1.4 倍為上限),使短欄獲較寬空間、減少換行
    avg = (sum(measure) / ncol) if ncol else 1
    capped = [min(c, 1.4 * avg) for c in measure]
    total = sum(capped) or 1
    TOTAL_CM = 15.6
    base = 2.6
    extra = max(0.0, TOTAL_CM - base * ncol)
    widths = [base + extra * (c / total) for c in capped]
    # 保證「label 欄」(內容 ≤22 半形之短標籤欄,非長文欄)能單行容納其最長內容,
    # 避免如表 1.1「RQ3 降幻覺與忠實引用」把「用」擠到次行;所需額外寬由最長(長文)欄吸收。
    if ncol:
        longest = max(range(ncol), key=lambda i: content_max[i])
        for i in range(ncol):
            if i != longest and content_max[i] <= 22:
                need = content_max[i] * 0.19 + 0.42  # 單行所需 cm(半形單位×字寬 + 左右邊距)
                if need > widths[i]:
                    widths[longest] = max(base, widths[longest] - (need - widths[i]))
                    widths[i] = need
    # 表頭:置中、粗體
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        cell.width = Cm(widths[i])
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h.strip())
        _set_run_fonts(r, size_pt=11, bold=True)
    # 內文:短內容(數字/✓✗/部分/短詞)置中,長文字靠左;固定欄寬使其自動換行
    # 整欄統一對齊:該欄所有內容皆短→整欄置中;含長文字→整欄靠左(避免同欄忽左忽右)
    col_center = []
    for ci in range(ncol):
        mx = 0
        for row in rows:
            if ci < len(row):
                mx = max(mx, _cell_wlen(row[ci]))
        col_center.append(mx <= 16)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            if ci >= ncol:
                continue
            cell = t.cell(ri, ci)
            cell.width = Cm(widths[ci])
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if col_center[ci]
                           else WD_ALIGN_PARAGRAPH.LEFT)
            _emit_inline(p, val.strip(), size_pt=10.5, code_size=9.5)
    # 長表跨頁時於次頁自動重複表頭列(NUTC 三線表慣例)
    _trPr = t.rows[0]._tr.get_or_add_trPr()
    _th = OxmlElement("w:tblHeader")
    _th.set(qn("w:val"), "true")
    _trPr.append(_th)
    _apply_three_line_table(t)
    # 表格與其後正文之間留間距(避免下一段正文緊貼表格底線);
    # 以小字級空段控制高度,約一個 6pt 行高的呼吸空間,而非整行空白。
    _sp = doc.add_paragraph()
    _sp.paragraph_format.space_before = Pt(0)
    _sp.paragraph_format.space_after = Pt(0)
    _sp.paragraph_format.line_spacing = 1.0
    _set_run_fonts(_sp.add_run(""), size_pt=6)


_TABLE_CAPTION = re.compile(r"^\s*\*{0,2}\s*表\s*\d+(?:\.[0-9A-Za-z]+)?\s*[:：]")


def _clean_caption(text: str) -> str:
    """去除 markdown 粗體標記,回傳乾淨之圖說/表題文字(供目次條目用)。"""
    return re.sub(r"\*\*", "", text).strip()


def _add_tc_field(paragraph, entry: str, flag: str) -> None:
    """在段落尾端附加隱藏 TC 欄位,供表目次(\\f t)/圖目次(\\f f)自動收集(含頁碼)。

    Word 開檔更新功能變數後,build_lot/build_lof 之 TOC \\f 欄位即據此列出條目與頁碼。
    """
    entry = entry.replace('"', "'").strip()
    if not entry:
        return
    r = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TC "{entry}" \\f {flag} \\l 1 '
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._element.append(begin)
    r._element.append(instr)
    r._element.append(end)


def _insert_landscape_figure(doc: DocxDocument, full_path: str, alt: str) -> None:
    """寬幅架構圖以「橫向滿版頁」呈現(台灣論文對寬圖之慣例),使文字夠大可辨識。
    前一段(直式)內容於此結束,圖置於 landscape section,圖後再轉回直式。"""
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    from PIL import Image as _PILImage
    try:
        with _PILImage.open(full_path) as _im:
            _iw, _ih = _im.size
    except Exception:
        _iw = _ih = 0
    # 橫向可用:寬 ~24.7cm、高 ~16cm;以寬 24cm 為主,過高則改高度上限 15.5cm
    if _iw and _ih and (_ih * 24.0 / _iw) > 15.5:
        p.add_run().add_picture(full_path, height=Cm(15.5))
    else:
        p.add_run().add_picture(full_path, width=Cm(24.0))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cr = cap.add_run(alt)
    _set_run_fonts(cr, size_pt=11)
    _add_tc_field(cap, _clean_caption(alt), "f")

    # 轉回直式
    port = doc.add_section(WD_SECTION_START.NEW_PAGE)
    port.orientation = WD_ORIENT.PORTRAIT
    port.page_width = Cm(21.0)
    port.page_height = Cm(29.7)
    port.top_margin = Cm(3)
    port.bottom_margin = Cm(3)
    port.left_margin = Cm(3)
    port.right_margin = Cm(3)


def render_markdown_file(doc: DocxDocument, md_path: Path) -> None:
    """把單一 markdown 檔案的內容 emit 到 docx。"""
    raw = md_path.read_text(encoding="utf-8")
    lines = raw.splitlines()
    in_code = False
    in_frontmatter = False
    code_buf: list[str] = []
    code_lang = ""
    # 跳過第一行的 H1 (重複章標題,因為章標題已由我們自加)
    skip_first_h1 = True
    skip_yaml_frontmatter = True
    skip_intro_quote = True  # 跳過開頭 > 引文 (作為章首備註)
    i = 0
    while i < len(lines):
        line = lines[i]
        # YAML frontmatter
        if i == 0 and _FRONTMATTER.match(line) and skip_yaml_frontmatter:
            in_frontmatter = True
            i += 1
            continue
        if in_frontmatter:
            if _FRONTMATTER.match(line):
                in_frontmatter = False
            i += 1
            continue
        # code fence
        if _CODE_FENCE.match(line):
            if in_code:
                emit_code_block(doc, code_buf, code_lang)
                code_buf = []
                in_code = False
            else:
                in_code = True
                code_lang = line.strip()[3:].strip()
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        # display math ($$...$$ 單行)
        m = _DISPLAY_MATH.match(line)
        if m:
            emit_math_block(doc, m.group(1).strip())
            i += 1
            continue
        # heading
        m = _HEADING.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if skip_first_h1 and level == 1:
                skip_first_h1 = False
                i += 1
                continue
            emit_heading(doc, min(level, 3), text)
            i += 1
            continue
        # image (SVG → 自動換成同名 PNG,因 python-docx 不接受 SVG)
        m = _IMG.search(line)
        if m:
            alt, path = m.group(1), m.group(2)
            full = (md_path.parent / path).resolve()
            if full.suffix.lower() == ".svg":
                png_alt = full.with_suffix(".png")
                if png_alt.exists():
                    full = png_alt
            if full.exists() and full.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".bmp"}:
                # 圖 3.1 系統架構(直向橫帶版):以接近滿版寬度直向插入(不再橫頁)
                if full.name == "fig_architecture.png":
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(0)
                        from PIL import Image as _PILImage
                        with _PILImage.open(str(full)) as _im:
                            _iw, _ih = _im.size
                        if _iw and (_ih * 15.0 / _iw) > 21.0:
                            p.add_run().add_picture(str(full), height=Cm(21.0))
                        else:
                            p.add_run().add_picture(str(full), width=Cm(15.0))
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap.paragraph_format.space_after = Pt(8)
                        cr = cap.add_run(alt)
                        _set_run_fonts(cr, size_pt=11)
                        _add_tc_field(cap, _clean_caption(alt), "f")
                    except Exception as e:  # noqa: BLE001
                        emit_paragraph(doc, f"[圖:{alt}]({path}) — 插入失敗: {e}")
                    i += 1
                    continue
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(0)
                    # 圖片以寬 14cm 為主;但過高之圖(如整頁 UI 截圖)改以高度上限
                    # 縮放,避免超出 A4 可用高度(29.7−3−3≈23.7cm)而被 Word 截斷。
                    from PIL import Image as _PILImage
                    _MAX_H_CM = 21.0
                    try:
                        with _PILImage.open(str(full)) as _im:
                            _iw, _ih = _im.size
                    except Exception:
                        _iw = _ih = 0
                    if _iw and _ih and (_ih * 14.0 / _iw) > _MAX_H_CM:
                        p.add_run().add_picture(str(full), height=Cm(_MAX_H_CM))
                    else:
                        p.add_run().add_picture(str(full), width=Cm(14))
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.paragraph_format.space_after = Pt(8)
                    cr = cap.add_run(alt)
                    _set_run_fonts(cr, size_pt=11)
                    _add_tc_field(cap, _clean_caption(alt), "f")  # 圖目次來源
                except Exception as e:
                    # 圖片插入失敗,只放圖說當佔位
                    emit_paragraph(doc, f"[圖:{alt}]({path}) — 插入失敗: {e}")
            else:
                emit_paragraph(doc, f"[圖:{alt}]({path})")
            i += 1
            continue
        # blockquote (skip first one as intro note)
        m = _BLOCKQUOTE.match(line)
        if m and skip_intro_quote:
            # consume consecutive blockquote lines
            while i < len(lines) and _BLOCKQUOTE.match(lines[i]):
                i += 1
            skip_intro_quote = False
            continue
        # table
        if line.strip().startswith("|") and i + 1 < len(lines) and _TABLE_SEP.match(lines[i + 1]):
            headers = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip header + separator
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            emit_table(doc, headers, rows)
            continue
        # list
        m = _LIST_BULLET.match(line)
        if m:
            emit_list_item(doc, m.group(1))
            i += 1
            continue
        m = _LIST_NUMBERED.match(line)
        if m:
            emit_list_item(doc, m.group(2), numbered=True, number=m.group(1))
            i += 1
            continue
        # 縮排字母子項 (a. b. c. …)：巢狀於數字項之下,左縮排加深
        m = _LIST_LETTER.match(line)
        if m:
            emit_list_item(
                doc, m.group(2), numbered=True, number=m.group(1), indent_base=1.48
            )
            i += 1
            continue
        # blockquote (general)
        m = _BLOCKQUOTE.match(line)
        if m:
            emit_paragraph(doc, m.group(1))
            i += 1
            continue
        # blank line
        if not line.strip():
            i += 1
            continue
        # 普通段落 (合併到下個空行)
        buf = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not (
            _HEADING.match(lines[i])
            or _LIST_BULLET.match(lines[i])
            or _LIST_NUMBERED.match(lines[i])
            or _LIST_LETTER.match(lines[i])
            or _CODE_FENCE.match(lines[i])
            or _BLOCKQUOTE.match(lines[i])
            or lines[i].strip().startswith("|")
        ):
            buf.append(lines[i])
            i += 1
        _txt = " ".join(buf)
        _para = emit_paragraph(doc, _txt)
        if _para is not None and _TABLE_CAPTION.match(_txt):
            _add_tc_field(_para, _clean_caption(_txt), "t")  # 表目次來源
            # 表題與其後之表格同頁(避免標題+表頭孤立於頁底、表身跑至次頁)
            _para.paragraph_format.keep_with_next = True
            _para.paragraph_format.keep_together = True
            # 表題與上文之間留間距(避免緊貼上一段正文)
            _para.paragraph_format.space_before = Pt(12)


# ────────────────────────── 各段落組裝 ──────────────────────────


def build_cover(doc: DocxDocument) -> None:
    """封面 (附件一)。"""
    add_blank_line(doc, 2)
    add_centered_paragraph(doc, "國立臺中科技大學資訊工程系碩士班", size_pt=18, cjk="標楷體-繁")
    add_blank_line(doc, 1)
    add_centered_paragraph(doc, "碩士論文", size_pt=22, bold=True, cjk="標楷體-繁")
    add_blank_line(doc, 3)
    add_centered_paragraph(doc, META["title_zh"], size_pt=24, bold=True, cjk="標楷體-繁")
    add_blank_line(doc, 1)
    add_centered_paragraph(doc, META["title_en"], size_pt=20, bold=True, cjk="Times New Roman")
    add_blank_line(doc, 4)
    # 掛名:左對齊 + 整塊縮排視覺置中(indent≈4cm)。前導 5 個全形空格(=「指導教授：」寬度)於左對齊下方能對齊第二位教授。
    _cov_adv = []
    for i, adv in enumerate(META["advisors"]):
        prefix = "指導教授：" if i == 0 else "　　　　　"
        _cov_adv.append(f"{prefix}{adv['zh']}    {adv['en']}")
    add_stacked_left_block(doc, _cov_adv, size_pt=14, cjk="標楷體-繁", indent=4.0)
    add_blank_line(doc, 1)
    add_left_paragraph(doc, f"研究生：{META['student_zh']}    {META['student_en']}", size_pt=14, cjk="標楷體-繁", indent=4.0)
    add_blank_line(doc, 3)
    add_centered_paragraph(
        doc, f"中華民國 {META['year_roc']} 年 {META['month_zh']} 月", size_pt=16, bold=True, cjk="標楷體-繁"
    )
    add_page_break(doc)


def build_approval(doc: DocxDocument) -> None:
    """學位考試委員會審定書 (附件三):置於書名頁之後、封面節(section1)內,不編頁碼。
    掃描影像置中滿版(寬 15 cm),浮水印在後被影像覆蓋。115/08/01 起強制納編。"""
    if not APPROVAL_PATH.exists():
        print(f"[!] 找不到審定書影像 {APPROVAL_PATH},略過")
        return
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(APPROVAL_PATH), width=Cm(15))


def build_title_page(doc: DocxDocument) -> None:
    """書名頁 (附件二)。"""
    add_blank_line(doc, 1)
    add_centered_paragraph(doc, META["title_zh"], size_pt=18, bold=True, cjk="標楷體-繁")
    add_blank_line(doc, 1)
    add_centered_paragraph(doc, META["title_en"], size_pt=14, bold=True, cjk="Times New Roman")
    add_blank_line(doc, 2)
    # 掛名:左對齊 + 整塊縮排視覺置中(同封面),使第二位指導教授對齊第一位正下方。
    _tp_adv = []
    for i, adv in enumerate(META["advisors"]):
        prefix = "指導教授：" if i == 0 else "　　　　　"
        _tp_adv.append(f"{prefix}{adv['zh']}    {adv['en']}")
    add_stacked_left_block(doc, _tp_adv, size_pt=14, cjk="標楷體-繁", indent=4.0)
    add_blank_line(doc, 1)
    add_left_paragraph(doc, f"研究生：{META['student_zh']}    {META['student_en']}", size_pt=14, cjk="標楷體-繁", indent=4.0)
    add_blank_line(doc, 2)
    add_centered_paragraph(doc, "國立臺中科技大學", size_pt=14, cjk="標楷體-繁")
    add_centered_paragraph(doc, "資訊工程系碩士班", size_pt=14, cjk="標楷體-繁")
    add_centered_paragraph(doc, "碩士論文", size_pt=14, cjk="標楷體-繁")
    add_blank_line(doc, 1)
    add_centered_paragraph(doc, "A Thesis", size_pt=12, cjk="Times New Roman")
    add_centered_paragraph(doc, "Submitted to", size_pt=12, cjk="Times New Roman")
    add_centered_paragraph(doc, META["dept_en"], size_pt=12, cjk="Times New Roman")
    add_centered_paragraph(doc, META["university_en"], size_pt=12, cjk="Times New Roman")
    add_centered_paragraph(doc, "in Partial Fulfillment of the Requirements", size_pt=12, cjk="Times New Roman")
    add_centered_paragraph(doc, "for the Degree of", size_pt=12, cjk="Times New Roman")
    add_centered_paragraph(doc, META["degree_en"], size_pt=12, cjk="Times New Roman")
    add_blank_line(doc, 2)
    add_centered_paragraph(doc, META["month_en"], size_pt=12, cjk="Times New Roman")
    add_centered_paragraph(doc, "Taichung, Taiwan, Republic of China", size_pt=12, cjk="Times New Roman")
    add_centered_paragraph(doc, f"中華民國 {META['year_roc']} 年 {META['month_zh']} 月", size_pt=14, bold=True, cjk="標楷體-繁")
    # 不在此手動換頁:下一步即為分節換頁(section break 到前置區),避免「換頁+分節」疊出空白頁


def build_zh_abstract(doc: DocxDocument) -> None:
    add_blank_line(doc, 1)
    add_centered_paragraph(doc, META["title_zh"], size_pt=16, bold=True)
    add_blank_line(doc, 1)
    # 掛名:左對齊 + 整塊縮排視覺置中(indent≈3.5cm)。續行前導 13 個全形空格(=「學生：許紫晴　　指導教授：」寬度)
    # 於左對齊下方能生效,使第二位教授「馬豪尚」對齊至第一位「張家瑋」正下方。
    _adv_lines = [f"學生：{META['student_zh']}　　指導教授：{META['advisors'][0]['zh']}"]
    for adv in META["advisors"][1:]:
        _adv_lines.append(f"　　　　　　　　　　　　　{adv['zh']}")
    add_stacked_left_block(doc, _adv_lines, size_pt=14, indent=3.5)
    add_blank_line(doc, 1)
    add_centered_paragraph(doc, "國立臺中科技大學資訊工程系碩士班", size_pt=14)
    add_blank_line(doc, 1)
    _add_tc_field(add_centered_paragraph(doc, "摘 要", size_pt=16, bold=True), "摘要", "c")
    add_blank_line(doc, 1)
    # body from 00_前置/中文摘要.md
    md = THESIS_DIR / "00_前置" / "中文摘要.md"
    if md.exists():
        for block in _read_abstract_body(md):
            if block.startswith("關鍵") or block[:8].lower().startswith("keyword"):
                add_blank_line(doc, 1)
            emit_paragraph(doc, block, size_pt=12)
    add_page_break(doc)


def build_en_abstract(doc: DocxDocument) -> None:
    add_blank_line(doc, 1)
    add_centered_paragraph(doc, META["title_en"], size_pt=16, bold=True, cjk="Times New Roman")
    add_blank_line(doc, 1)
    # 掛名:左對齊 + 整塊縮排視覺置中。第二位指導教授以 Tab stop 對齊至第一位正下方
    # (比例字體 Times New Roman 下,固定空格寬度對不準會使第二位教授飄移;改用 Tab
    #  定位則與字體無關皆精準對齊)。Tab 位置 = 縮排 + 前綴實測寬度(PIL 量測,附 fallback)。
    _en_indent = 3.5
    _prefix_en = f"Student: {META['student_en']}    Advisor: "
    _tab_cm = _en_indent + 7.6  # fallback:前綴寬估值(字型缺失時)
    try:
        from PIL import ImageFont as _ImageFont
        _fnt = _ImageFont.truetype(
            "/System/Library/Fonts/Supplemental/Times New Roman.ttf", 14)
        _tab_cm = _en_indent + _fnt.getlength(_prefix_en) / 72.0 * 2.54 + 0.1
    except Exception:  # noqa: BLE001
        pass
    _pa = doc.add_paragraph()
    _pa.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _paf = _pa.paragraph_format
    _paf.space_before = Pt(0)
    _paf.space_after = Pt(0)
    _paf.line_spacing = 1.0
    _paf.left_indent = Cm(_en_indent)
    _paf.tab_stops.add_tab_stop(Cm(_tab_cm), WD_TAB_ALIGNMENT.LEFT)
    _r = _pa.add_run(f"Student: {META['student_en']}    Advisor:")
    _set_run_fonts(_r, cjk="Times New Roman", size_pt=14)
    _r = _pa.add_run("\t" + META["advisors"][0]["en"])
    _set_run_fonts(_r, cjk="Times New Roman", size_pt=14)
    for adv in META["advisors"][1:]:
        _pa.add_run().add_break()
        _r = _pa.add_run("\t" + adv["en"])
        _set_run_fonts(_r, cjk="Times New Roman", size_pt=14)
    add_blank_line(doc, 1)
    add_centered_paragraph(doc, META["dept_en"], size_pt=14, cjk="Times New Roman")
    add_centered_paragraph(doc, META["university_en"], size_pt=14, cjk="Times New Roman")
    add_blank_line(doc, 1)
    _add_tc_field(add_centered_paragraph(doc, "ABSTRACT", size_pt=16, bold=True, cjk="Times New Roman"), "ABSTRACT", "c")
    add_blank_line(doc, 1)
    md = THESIS_DIR / "00_前置" / "Abstract.md"
    if md.exists():
        for block in _read_abstract_body(md):
            if block.startswith("關鍵") or block[:8].lower().startswith("keyword"):
                add_blank_line(doc, 1)
            emit_paragraph(doc, block, size_pt=12)
    add_page_break(doc)


def _read_abstract_body(md: Path) -> list[str]:
    """讀摘要 markdown,跳過第一個 H1 (重複標題),回傳每段純文字。"""
    lines = md.read_text(encoding="utf-8").splitlines()
    blocks: list[str] = []
    buf: list[str] = []
    seen_h1 = False
    for line in lines:
        if line.startswith("# "):
            seen_h1 = True
            continue
        if not seen_h1:
            continue
        if not line.strip():
            if buf:
                blocks.append(" ".join(buf))
                buf = []
            continue
        buf.append(line.strip())
    if buf:
        blocks.append(" ".join(buf))
    return blocks


def build_acknowledgements(doc: DocxDocument) -> None:
    """誌謝(編印順序:英摘之後、目次之前)。內容讀自 00_前置/致謝.md,首段首行對齊正文。"""
    md = THESIS_DIR / "00_前置" / "致謝.md"
    if not md.exists():
        return
    add_blank_line(doc, 1)
    _add_tc_field(add_centered_paragraph(doc, "誌　謝", size_pt=16, bold=True), "誌謝", "c")
    add_blank_line(doc, 1)
    blocks = _read_abstract_body(md)
    # 署名與日期(含「謹誌」或「中華民國…年…月」)靠右對齊,且與正文間留一空行
    _SIGN_RE = re.compile(r"謹誌|中華民國.*年.*月")
    sign_started = False
    for block in blocks:
        is_sign = bool(_SIGN_RE.search(block))
        if is_sign and not sign_started:
            add_blank_line(doc, 1)   # 正文與署名間留白
            sign_started = True
        p = emit_paragraph(doc, block, size_pt=12)
        if p is not None:
            if is_sign:
                # 署名/日期:靠右、上下黏緊(兩行相鄰不留間距)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.first_line_indent = Cm(0)
            else:
                p.paragraph_format.space_after = Pt(8)
    # 不在此手動 add_page_break:當誌謝剛好寫滿整頁時,空的換頁段落會溢出成一整頁空白。
    # 改由目次標題之「段前分頁」達成換頁(見 build_toc),不留懸空空段落。


def build_toc(doc: DocxDocument) -> None:
    """目次 - 插入 Word 自動目次欄位 (使用者開檔後右鍵『更新欄位』即可)。"""
    _toc_title = add_centered_paragraph(doc, "目  次", size_pt=14, bold=True)
    _toc_title.paragraph_format.page_break_before = True  # 段前分頁:不留空白頁
    # 目次自身亦納入自動目次(連結+頁碼);前置各項(摘要/ABSTRACT/誌謝/表目次/圖目次)
    # 已於各自標題以 TC \f c 標記,故此目次自動含前置至圖目次全部頁碼與超連結。
    _add_tc_field(_toc_title, "目次", "c")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run()
    _set_run_fonts(r, size_pt=12)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\f c \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "(在 Word 內按 F9 或右鍵『更新功能變數』以產生目次)"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r._element.append(fld_begin)
    r._element.append(instr)
    r._element.append(fld_sep)
    r._element.append(t)
    r._element.append(fld_end)
    # 不在此手動換頁:表目次以段前分頁起始,避免目次剛好填滿頁時多擠出一頁空白。


def _insert_toc_field(doc: DocxDocument, instr_text: str, placeholder: str) -> None:
    """插入一個 TOC 類功能變數(目次/表目次/圖目次共用結構)。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run()
    _set_run_fonts(r, size_pt=12)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, t, end):
        r._element.append(el)


def build_lot(doc: DocxDocument) -> None:
    """表目次(附件八)。由各表題之 TC \\f t 欄位填充(Word 更新功能變數後生效)。"""
    _t = add_centered_paragraph(doc, "表 目 次", size_pt=14, bold=True)
    _t.paragraph_format.page_break_before = True  # 各自獨立頁(規範附件八)
    _add_tc_field(_t, "表目次", "c")  # 納入主目次(連結+頁碼)
    add_blank_line(doc, 1)
    _insert_toc_field(
        doc, "TOC \\h \\z \\f t", "(在 Word 內按 F9 或右鍵『更新功能變數』以產生表目次)"
    )


def build_lof(doc: DocxDocument) -> None:
    """圖目次(附件九)。由各圖說之 TC \\f f 欄位填充(Word 更新功能變數後生效)。"""
    _t = add_centered_paragraph(doc, "圖 目 次", size_pt=14, bold=True)
    _add_tc_field(_t, "圖目次", "c")  # 納入主目次(連結+頁碼)
    _t.paragraph_format.page_break_before = True  # 各自獨立頁(規範附件九)
    add_blank_line(doc, 1)
    _insert_toc_field(
        doc, "TOC \\h \\z \\f f", "(在 Word 內按 F9 或右鍵『更新功能變數』以產生圖目次)"
    )
    # 不在此手動換頁:下一步即為分節換頁(section break 到本文),避免疊出空白頁


# ────────────────────────── 主流程 ──────────────────────────


def _no_overflow_punct_on_ppr(ppr) -> None:
    """在給定之 w:pPr 加入 <w:overflowPunct w:val="0"/>(關閉行尾標點溢出邊界)。"""
    op = ppr.find(qn("w:overflowPunct"))
    if op is None:
        op = OxmlElement("w:overflowPunct")
        ppr.insert(0, op)  # overflowPunct 於 CT_PPrBase 序位靠前,置首最安全
    op.set(qn("w:val"), "0")


def _set_docdefaults_no_overflow_punct(doc: DocxDocument) -> None:
    """於 styles.xml 之 docDefaults/pPrDefault/pPr 設 overflowPunct=0,作全域預設。"""
    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        dd = OxmlElement("w:docDefaults")
        styles_el.insert(0, dd)
    ppd = dd.find(qn("w:pPrDefault"))
    if ppd is None:
        ppd = OxmlElement("w:pPrDefault")
        dd.append(ppd)
    ppr = ppd.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        ppd.append(ppr)
    _no_overflow_punct_on_ppr(ppr)


def configure_default_style(doc: DocxDocument) -> None:
    """設定 Normal style 為新細明體 + Times,12 pt,單行間距。"""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "新細明體")

    # 關閉「允許標點溢出邊界」:行尾標點(,。、等)不再懸掛凸出右邊界,右緣乾淨對齊。
    # 設於 Normal 之 pPr(全文段落皆繼承)+ docDefaults,雙保險確保各式段落一致。
    _no_overflow_punct_on_ppr(style.element.get_or_add_pPr())
    _set_docdefaults_no_overflow_punct(doc)

    # heading styles
    for h, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        st = doc.styles[h]
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "新細明體")


def _set_update_fields_on_open(doc: DocxDocument) -> None:
    """令 Word 開檔時自動更新所有功能變數(目次/表目次/圖目次/頁碼),
    使轉出之 PDF 直接呈現已填充之各式目次,毋須手動 F9。"""
    try:
        settings = doc.settings.element
        if settings.find(qn("w:updateFields")) is None:
            el = OxmlElement("w:updateFields")
            el.set(qn("w:val"), "true")
            settings.append(el)
    except Exception:  # noqa: BLE001
        pass


def main() -> int:
    _EQ_NO[0] = 0
    doc = Document()
    configure_default_style(doc)
    tighten_toc_styles(doc)

    # ── Section 1: 封面 + 書名頁 (無頁碼) ──
    section1 = doc.sections[0]
    _setup_section_page_format(section1, page_num_fmt="decimal", start=1)
    # 封面 / 書名頁不顯示頁碼 → footer 留空
    section1.footer.is_linked_to_previous = False
    for p in list(section1.footer.paragraphs):
        p_elem = p._element
        # 留一個空段落,不放欄位
        p.clear()
    build_cover(doc)
    build_title_page(doc)
    build_approval(doc)  # 學位考試委員會審定書 (書名頁後,不編頁碼)

    # ── Section 2: 前置 (摘要 → 目次),小寫羅馬,從 i 開始 ──
    section2 = add_section_break(doc)
    _setup_section_page_format(
        section2, page_num_fmt="lowerRoman", start=1, restart=True
    )
    _add_footer_page_number(section2)
    build_zh_abstract(doc)
    build_en_abstract(doc)
    build_acknowledgements(doc)
    build_toc(doc)
    build_lot(doc)
    build_lof(doc)

    # ── Section 3: 論文本文 (§1–§7) + 參考文獻 + 附錄,阿拉伯,從 1 開始 ──
    section3 = add_section_break(doc)
    _setup_section_page_format(
        section3, page_num_fmt="decimal", start=1, restart=True
    )
    _add_footer_page_number(section3)

    chapters = [
        ("01_緒論.md", "第一章 緒論"),
        ("02_文獻探討.md", "第二章 文獻探討"),
        ("03_系統設計.md", "第三章 系統設計"),
        ("04_實作細節.md", "第四章 實作細節"),
        ("05_延伸機制.md", "第五章 延伸機制"),
        ("06_實驗與結果.md", "第六章 實驗與結果"),
        ("07_結論.md", "第七章 結論與未來工作"),
    ]
    # 每章以「標題前換頁」起新頁(第一章除外:本文 section 已自起新頁),
    # 取代尾端 add_page_break,避免前一章寫到頁尾時溢出整頁空白(原 §3↔§4 之第 27 頁)
    for idx, (fname, ch_title) in enumerate(chapters):
        emit_heading(doc, 1, ch_title, page_break_before=(idx > 0))
        render_markdown_file(doc, THESIS_DIR / fname)

    # 參考文獻
    emit_heading(doc, 1, "參考文獻", page_break_before=True)
    render_markdown_file(doc, THESIS_DIR / "參考文獻.md")

    # 附錄
    appendix_dir = THESIS_DIR / "附錄"
    if appendix_dir.exists():
        for app_md in sorted(appendix_dir.glob("*.md")):
            name = app_md.stem  # e.g. "A_System_Prompts"
            letter, _, rest = name.partition("_")
            title = f"附錄 {letter} {rest.replace('_', ' ')}".strip()
            emit_heading(doc, 1, title, page_break_before=True)
            render_markdown_file(doc, app_md)

    _set_update_fields_on_open(doc)
    doc.save(OUT_PATH)

    if WATERMARK_PATH.exists():
        add_watermark_post(OUT_PATH, WATERMARK_PATH)
        print(f"[OK] 已加入浮水印 ({WATERMARK_PATH.name})")
    else:
        print(f"[!] 找不到浮水印圖檔 {WATERMARK_PATH}")

    print(f"[OK] 寫入 {OUT_PATH}")
    print(f"     大小: {OUT_PATH.stat().st_size // 1024} KB")

    # ── 同步輸出 PDF ──
    # 改用 AppleScript 直接驅動 Word:`display alerts none` 自動確認「更新功能變數」對話框
    # (目次/表目次/圖目次/頁碼自動填入)、`with timeout 600` 避免大檔逾時,較 docx2pdf 穩定。
    import subprocess
    pdf_path = OUT_PATH.with_suffix(".pdf")
    if pdf_path.exists():
        try:
            pdf_path.unlink()
        except OSError:
            pass
    # 先關閉「開檔自動更新」旗標,避免 Word 開檔時跳出「是否更新功能變數?」modal 阻塞
    # AppleScript(此為大檔常見卡死主因;Word 0% CPU、無視窗)。目次改由下方 AppleScript
    # 以 `update (tables of contents)` 明確重算(不受此旗標影響),仍會產生含 PAGEREF 之可點目次。
    try:
        import zipfile as _zf
        import re as _re2
        with _zf.ZipFile(OUT_PATH) as _z:
            _names = _z.namelist()
            _data = {n: _z.read(n) for n in _names}
        if "word/settings.xml" in _data:
            _s = _data["word/settings.xml"].decode("utf-8")
            _s = _re2.sub(r'<w:updateFields[^/]*/>', '<w:updateFields w:val="false"/>', _s)
            _data["word/settings.xml"] = _s.encode("utf-8")
            with _zf.ZipFile(OUT_PATH, "w", _zf.ZIP_DEFLATED) as _zo:
                for n in _names:
                    _zo.writestr(n, _data[n])
    except Exception as _e:  # noqa: BLE001
        print(f"[!] 預先關閉 updateFields 略過:{str(_e)[:60]}")

    # AppleScript:只更新目次/圖表目次物件(數個,快速且穩定),再存回 docx 與輸出 PDF。
    # 逐一 update 每個 field(數百個)於大檔會極慢,故只更新 tables of contents/figures。
    # delay 需足夠讓大檔(120+ 頁)完全載入分頁,否則 update 於未載完狀態產生空目次。
    # 以「索引存取 table of contents i」較 repeat-in 穩定(跨呼叫 active document 較不易失效)。
    _as = f'''with timeout of 1800 seconds
    tell application "Microsoft Word"
        activate
        set display alerts to none
        open POSIX file "{OUT_PATH}"
        delay 6
        set d to active document
        set n to (count of (get tables of contents of d))
        repeat with i from 1 to n
            update (table of contents i of d)
        end repeat
        try
            set m to (count of (get tables of figures of d))
            repeat with i from 1 to m
                update (table of figures i of d)
            end repeat
        end try
        save d
        save as d file name "{pdf_path}" file format format PDF
        close d saving no
        quit saving no
    end tell
end timeout'''
    word_ok = False
    # Word 自動化偶發失敗(目次未更新→顯示「按 F9」提示)。重試一次;每次前先清 Word 狀態。
    for _attempt in range(2):
        try:
            subprocess.run(
                ["osascript", "-e", 'tell application "Microsoft Word" to quit saving no'],
                capture_output=True, text=True, timeout=60)
        except Exception:  # noqa: BLE001
            pass
        try:
            subprocess.run(["osascript", "-e", _as], check=True,
                           capture_output=True, text=True, timeout=1860)
            word_ok = True
            break
        except Exception as e:  # noqa: BLE001
            print(f"[!] AppleScript 第 {_attempt + 1} 次失敗:{str(e)[:60]}")
    if not word_ok:
        print("[!] AppleScript 兩次皆失敗,改試 docx2pdf(目次不會自動生成)")
        try:
            from docx2pdf import convert
            convert(str(OUT_PATH), str(pdf_path))
        except Exception as e2:  # noqa: BLE001
            print(f"[!] PDF 轉換略過(需 Word):{str(e2)[:80]}")
    if pdf_path.exists():
        print(f"[OK] 同步輸出 PDF {pdf_path}")
        print(f"     大小: {pdf_path.stat().st_size // 1024} KB")
    # 僅在 Word 成功更新+存檔(docx 已含快取目次)時才收尾:關閉自動更新提示 + 清 metadata。
    # 若 Word 未執行(無 Word),保留 updateFields=true,使目次仍能於開檔時生成。
    if word_ok:
        _collapse_toc_trailing_blank(OUT_PATH)
        _finalize_docx(OUT_PATH)
    return 0


def _collapse_toc_trailing_blank(docx_path: Path) -> None:
    """消除「目次後空白頁」:Word 生成長目次後,欄位結尾常留一個空段落;當其緊接
    『段前分頁』之標題(表目次/圖目次/第一章)時,該空段落會被擠上獨立一頁而形成空白頁。
    將此空段落之 runs(含欄位結尾)併入前一段並刪除之,消除空白頁而不動搖已快取之目次。"""
    from docx import Document as _Doc
    try:
        d = _Doc(str(docx_path))
        ps = d.paragraphs
        removed = 0
        for i in range(1, len(ps) - 1):
            p = ps[i]
            if p.text.strip():
                continue
            if not p._p.findall(".//" + qn("w:fldChar")):
                continue  # 僅處理裝著欄位結尾之空段落
            if not ps[i + 1].paragraph_format.page_break_before:
                continue  # 僅在其後為段前分頁之標題時才會造成空白頁
            prev = ps[i - 1]._p
            for r in list(p._p.findall(qn("w:r"))):
                p._p.remove(r)
                prev.append(r)
            p._p.getparent().remove(p._p)
            removed += 1
        if removed:
            d.save(str(docx_path))
            print(f"[OK] 消除目次後空白頁 x{removed}")
    except Exception as e:  # noqa: BLE001
        print(f"[!] 目次空白頁處理略過:{str(e)[:60]}")


def _finalize_docx(docx_path: Path) -> None:
    """交件前收尾:①移除「開檔自動更新功能變數」旗標(避免開檔跳對話框;目次已快取)
    ②清空個資 metadata(作者設為研究生姓名、清最後儲存者/公司/描述)。純 zip 操作,不需 Word。"""
    import zipfile
    import re as _re
    try:
        with zipfile.ZipFile(docx_path) as zin:
            names = zin.namelist()
            data = {n: zin.read(n) for n in names}
        # ① 關閉開檔自動更新
        if "word/settings.xml" in data:
            s = data["word/settings.xml"].decode("utf-8")
            s = _re.sub(r'<w:updateFields[^/]*/>', '<w:updateFields w:val="false"/>', s)
            data["word/settings.xml"] = s.encode("utf-8")
        # ② 清 metadata
        if "docProps/core.xml" in data:
            c = data["docProps/core.xml"].decode("utf-8")
            c = _re.sub(r'(<dc:creator>)[^<]*(</dc:creator>)',
                        r'\g<1>' + META.get("student_zh", "") + r'\g<2>', c)
            c = _re.sub(r'<cp:lastModifiedBy>[^<]*</cp:lastModifiedBy>',
                        '<cp:lastModifiedBy></cp:lastModifiedBy>', c)
            c = _re.sub(r'(<dc:description>)[^<]*(</dc:description>)', r'\g<1>\g<2>', c)
            data["docProps/core.xml"] = c.encode("utf-8")
        if "docProps/app.xml" in data:
            a = data["docProps/app.xml"].decode("utf-8")
            a = _re.sub(r'(<Company>)[^<]*(</Company>)', r'\g<1>\g<2>', a)
            a = _re.sub(r'(<Manager>)[^<]*(</Manager>)', r'\g<1>\g<2>', a)
            data["docProps/app.xml"] = a.encode("utf-8")
        tmp = docx_path.with_suffix(".final.tmp")
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for n in names:
                zout.writestr(n, data[n])
        tmp.replace(docx_path)
        print(f"[OK] 交件收尾:已關閉開檔自動更新 + 清 metadata(作者={META.get('student_zh','')})")
    except Exception as e:  # noqa: BLE001
        print(f"[!] 交件收尾略過:{str(e)[:80]}")


# ────────────────────────── 浮水印 post-processing ──────────────────────────


def add_watermark_post(docx_path: Path, image_path: Path) -> None:
    """把指定 jpeg 嵌入 docx,作為所有 section 之 header VML watermark。

    流程:
      1. 加 word/media/watermark.jpeg
      2. 寫一份 word/header_watermark.xml (含 v:pict VML)
      3. 寫 word/_rels/header_watermark.xml.rels 指到 image
      4. 在 word/_rels/document.xml.rels 加 header relationship
      5. 在 document.xml 每個 <w:sectPr> 開頭插 <w:headerReference>
      6. 在 [Content_Types].xml 註冊 header_watermark.xml 與 jpeg
    """
    import re
    import shutil
    import zipfile

    src = Path(docx_path)
    tmp = src.with_suffix(".tmp.docx")

    with zipfile.ZipFile(src, "r") as zin:
        members = {n: zin.read(n) for n in zin.namelist()}

    # 1. 加圖
    members["word/media/watermark.jpeg"] = Path(image_path).read_bytes()

    # 2/4. 分配 rId
    rels_xml = members["word/_rels/document.xml.rels"].decode("utf-8")
    used = [int(x) for x in re.findall(r'Id="rId(\d+)"', rels_xml)]
    next_id = (max(used) if used else 0) + 1
    header_rid = f"rId{next_id}"
    image_rid_in_header = "rId1"  # header rels 內部獨立空間,從 rId1 起算

    # 3. header rels
    members["word/_rels/header_watermark.xml.rels"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        f'<Relationship Id="{image_rid_in_header}" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
        'Target="media/watermark.jpeg"/>'
        "</Relationships>"
    ).encode("utf-8")

    # 4. header XML (含 VML watermark)
    header_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word">'
        "<w:p><w:pPr><w:pStyle w:val=\"Header\"/></w:pPr>"
        "<w:r>"
        "<w:pict>"
        '<v:shapetype id="_x0000_t75" coordsize="21600,21600" o:spt="75" o:preferrelative="t" '
        'path="m@4@5l@4@11@9@11@9@5xe" filled="f" stroked="f">'
        '<v:stroke joinstyle="miter"/>'
        '<v:formulas>'
        '<v:f eqn="if lineDrawn pixelLineWidth 0"/>'
        '<v:f eqn="sum @0 1 0"/>'
        '<v:f eqn="sum 0 0 @1"/>'
        '<v:f eqn="prod @2 1 2"/>'
        '<v:f eqn="prod @3 21600 pixelWidth"/>'
        '<v:f eqn="prod @3 21600 pixelHeight"/>'
        '<v:f eqn="sum @0 0 1"/>'
        '<v:f eqn="prod @6 1 2"/>'
        '<v:f eqn="prod @7 21600 pixelWidth"/>'
        '<v:f eqn="sum @8 21600 0"/>'
        '<v:f eqn="prod @7 21600 pixelHeight"/>'
        '<v:f eqn="sum @10 21600 0"/>'
        '</v:formulas>'
        '<v:path o:extrusionok="f" gradientshapeok="t" o:connecttype="rect"/>'
        '<o:lock v:ext="edit" aspectratio="t"/>'
        '</v:shapetype>'
        '<v:shape id="WordPictureWatermark1" o:spid="_x0000_s2049" type="#_x0000_t75" '
        'style="position:absolute;left:0;text-align:left;margin-left:0;margin-top:0;'
        'width:153pt;height:156pt;z-index:-251657216;'
        'mso-position-horizontal:center;mso-position-horizontal-relative:margin;'
        'mso-position-vertical:center;mso-position-vertical-relative:margin" '
        'o:allowincell="f">'
        f'<v:imagedata r:id="{image_rid_in_header}" o:title="watermark"/>'
        '<w10:wrap anchorx="margin" anchory="margin"/>'
        '</v:shape>'
        "</w:pict>"
        "</w:r>"
        "</w:p>"
        "</w:hdr>"
    )
    members["word/header_watermark.xml"] = header_xml.encode("utf-8")

    # 5. 在 document.xml.rels 加 header relationship
    new_rel = (
        f'<Relationship Id="{header_rid}" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/header" '
        'Target="header_watermark.xml"/>'
    )
    rels_xml = rels_xml.replace("</Relationships>", new_rel + "</Relationships>")
    members["word/_rels/document.xml.rels"] = rels_xml.encode("utf-8")

    # 6. 每個 <w:sectPr...> 開頭插 headerReference (避免重複插入)
    doc_xml = members["word/document.xml"].decode("utf-8")
    header_ref_tag = (
        f'<w:headerReference xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        f'w:type="default" r:id="{header_rid}"/>'
    )

    def _inject(match):
        opening = match.group(0)
        # 把 headerReference 插在 opening tag 之後
        return opening + header_ref_tag

    doc_xml = re.sub(r"<w:sectPr[^/]*?>", _inject, doc_xml)
    # 規範要求「所有頁面」皆插浮水印,故電子版封面亦保留浮水印(不設 titlePg 例外)。
    # 列印版之封面浮水印移除,由 make_spine_print.py 另行處理。
    members["word/document.xml"] = doc_xml.encode("utf-8")

    # 7. [Content_Types].xml 註冊新檔
    ct_xml = members["[Content_Types].xml"].decode("utf-8")
    if "header_watermark.xml" not in ct_xml:
        override = (
            '<Override PartName="/word/header_watermark.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"/>'
        )
        ct_xml = ct_xml.replace("</Types>", override + "</Types>")
    # 確保 jpeg 副檔名有 Content-Type
    if 'Extension="jpeg"' not in ct_xml and 'Extension="jpg"' not in ct_xml:
        ct_xml = ct_xml.replace(
            "<Types ",
            '<Types ',
        )
        # 在第一個 Default 後加 jpeg default
        ct_xml = re.sub(
            r"(<Default[^/]*/>)",
            r'\1<Default Extension="jpeg" ContentType="image/jpeg"/>',
            ct_xml,
            count=1,
        )
    members["[Content_Types].xml"] = ct_xml.encode("utf-8")

    # 8. 寫回
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in members.items():
            zout.writestr(name, data)
    shutil.move(str(tmp), str(src))


if __name__ == "__main__":
    raise SystemExit(main())
